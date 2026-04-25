from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
TASK_DIR = ROOT / ".codex-tasks" / "20260425-thesis-docx"
RAW_DIR = TASK_DIR / "raw"
DOCX_PATH = ROOT / "NFT数字藏品交易平台论文.docx"
FIG_DIR = TASK_DIR / "figures"
TITLE = "NFT交易平台设计与实现"
BODY_FONT = "宋体"
LATIN_FONT = "Times New Roman"


def ensure_dirs() -> None:
    RAW_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    apply_page_layout(section)
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(20)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[style_name]
        style.font.name = BODY_FONT
        style.font.bold = True
        style.font.size = Pt(size)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        style.paragraph_format.line_spacing = Pt(20)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)

    doc.settings.odd_and_even_pages_header_footer = True


def apply_page_layout(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def set_run_font(run, size: int | float = 12, bold: bool = False, color: str | None = None) -> None:
    run.font.name = LATIN_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)


def add_paragraph(
    doc: Document,
    text: str = "",
    *,
    style: str | None = None,
    align: int | None = None,
    first_line: bool = True,
    size: int | float = 12,
    bold: bool = False,
    space_after: int | float = 0,
) -> None:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if not first_line:
        p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold)


def add_heading(doc: Document, text: str, level: int = 1, numbered: bool = True) -> None:
    p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, size=16 if level == 1 else 14 if level == 2 else 12, bold=True)


def set_static_footer(section, text: str) -> None:
    section.footer.is_linked_to_previous = False
    clear_container(section.footer)
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=10.5)


def clear_container(container) -> None:
    for p in container.paragraphs:
        for run in p.runs:
            run._element.getparent().remove(run._element)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char1)

    run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE \\* Arabic "
    run._r.append(instr)

    run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    run = paragraph.add_run()
    text = OxmlElement("w:t")
    text.text = "1"
    run._r.append(text)

    run = paragraph.add_run()
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char2)


def restart_page_numbering(section, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))
    pg_num.set(qn("w:fmt"), "decimal")


def setup_body_header_footer(section) -> None:
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = False
    section.even_page_header.is_linked_to_previous = False
    section.first_page_header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.even_page_footer.is_linked_to_previous = False
    section.first_page_footer.is_linked_to_previous = False
    clear_container(section.header)
    clear_container(section.even_page_header)
    clear_container(section.first_page_header)
    clear_container(section.footer)
    clear_container(section.even_page_footer)
    clear_container(section.first_page_footer)

    odd = section.header.paragraphs[0]
    odd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    odd.paragraph_format.first_line_indent = Pt(0)
    r = odd.add_run(TITLE)
    set_run_font(r, size=10.5)

    even = section.even_page_header.paragraphs[0]
    even.alignment = WD_ALIGN_PARAGRAPH.LEFT
    even.paragraph_format.first_line_indent = Pt(0)
    r = even.add_run("本科毕业设计（论文）")
    set_run_font(r, size=10.5)

    first = section.first_page_header.paragraphs[0]
    first.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    first.paragraph_format.first_line_indent = Pt(0)
    r = first.add_run(TITLE)
    set_run_font(r, size=10.5)

    for footer_part in [section.footer, section.even_page_footer, section.first_page_footer]:
        footer = footer_part.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.paragraph_format.first_line_indent = Pt(0)
        add_page_field(footer)


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    section.footer.is_linked_to_previous = False
    clear_container(section.footer)

    add_paragraph(doc, "\n\n", first_line=False)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run("毕业设计（论文）")
    set_run_font(r, size=22, bold=True)

    add_paragraph(doc, "\n", first_line=False)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(TITLE)
    set_run_font(r, size=18, bold=True)

    add_paragraph(doc, "\n\n\n", first_line=False)
    fields = [
        ("系    别", "____________________"),
        ("专    业", "____________________"),
        ("班    级", "____________________"),
        ("学生姓名", "____________________"),
        ("学    号", "____________________"),
        ("指导教师", "____________________"),
        ("完成日期", "2026 年 4 月"),
    ]
    for name, value in fields:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        r = p.add_run(f"{name}：{value}")
        set_run_font(r, size=18)


def abstract_text() -> list[str]:
    return [
        "随着数字内容生产和网络传播速度不断提升，图片、音频、视频等数字作品在复制、分发和二次传播中面临归属难确认、交易过程难追溯、创作者收益难延续等问题。区块链具有去中心化、难篡改和可追溯等特点，NFT 则能够为数字藏品建立唯一的链上标识，使作品从普通文件转变为可验证、可流转的数字资产。因此，围绕数字藏品的确权、展示和交易需求，设计并实现一个具备完整业务闭环的 NFT 交易平台，具有较强的工程实践意义。",
        "本文以 NFT 数字藏品交易场景为研究对象，采用前后端分离、链上链下协同的技术路线进行系统设计。前端基于 Next.js、React 与 ethers.js 实现市场浏览、钱包登录、作品创建、交易反馈和个人中心等功能；后端基于 Go、Gorilla Mux、GORM 与 MySQL 实现用户认证、NFT 数据管理、订单管理、文件校验和 IPFS 上传；智能合约基于 Solidity、ERC-721 与 EIP-2981 设计，支持铸造、上架、下架、购买和版税结算。系统最终实现了钱包签名登录、多媒体 NFT 创建、IPFS 元数据存储、链上交易、订单同步、资产管理和通知提醒等核心功能。",
        "实践结果表明，本系统能够较完整地支撑数字藏品从上传、铸造、展示、购买到个人资产管理的全过程。系统将确权、支付和版税等关键逻辑放在智能合约中执行，将高频查询、页面展示和用户资料管理放在链下服务中处理，在可信性与可用性之间取得了较好的平衡。该设计体现了区块链、NFT、智能合约、IPFS 与 Web 应用开发技术的综合应用价值，也为后续扩展事件监听、自动化测试和多链适配奠定了基础。",
    ]


def add_abstract(doc: Document) -> None:
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_page_layout(sec)
    set_static_footer(sec, "Ⅰ")

    add_paragraph(doc, TITLE, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=16, bold=True, space_after=8)
    add_paragraph(doc, "摘要", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=16, bold=True, space_after=8)
    for para in abstract_text():
        add_paragraph(doc, para, size=12)
    add_paragraph(doc, "", first_line=False, space_after=4)
    add_paragraph(doc, "关键词：区块链；NFT；智能合约；IPFS；数字藏品", first_line=False, size=12, bold=True)


def add_toc(doc: Document) -> None:
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_page_layout(sec)
    set_static_footer(sec, "Ⅱ")

    add_paragraph(doc, "目录", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=16, bold=True, space_after=14)
    entries = [
        ("摘要", "Ⅰ", 0),
        ("第1章 绪论", "1", 0),
        ("1.1 选题背景", "1", 1),
        ("1.2 研究目的与意义", "1", 1),
        ("1.3 主要研究内容", "1", 1),
        ("第2章 相关开发技术", "3", 0),
        ("第3章 系统分析", "5", 0),
        ("第4章 系统设计", "7", 0),
        ("第5章 系统实现", "11", 0),
        ("第6章 系统测试", "13", 0),
        ("结论", "15", 0),
        ("参考文献", "16", 0),
        ("附录", "17", 0),
        ("致谢", "18", 0),
    ]
    for title, page, indent in entries:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(18 * indent)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(20)
        r1 = p.add_run(title)
        set_run_font(r1, size=12)
        dots = "." * max(5, 48 - len(title) * 2 - indent * 4)
        r2 = p.add_run(dots)
        set_run_font(r2, size=12)
        r3 = p.add_run(page)
        set_run_font(r3, size=12)


def get_font(size: int, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def draw_centered(draw, box, text, font, fill="#111827") -> None:
    x1, y1, x2, y2 = box
    bbox = draw.textbbox((0, 0), text, font=font)
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.text((x1 + (x2 - x1 - w) / 2, y1 + (y2 - y1 - h) / 2), text, font=font, fill=fill)


def make_diagrams() -> tuple[Path, Path]:
    font = get_font(34)
    small = get_font(26)
    arch = FIG_DIR / "architecture.png"
    flow = FIG_DIR / "flow.png"

    img = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(img)
    boxes = [
        (90, 120, 360, 250, "前端表示层\nNext.js / React"),
        (500, 120, 780, 250, "后端服务层\nGo / REST API"),
        (920, 120, 1200, 250, "数据持久层\nMySQL / GORM"),
        (500, 420, 780, 550, "智能合约层\nERC-721 / 版税"),
        (920, 420, 1200, 550, "分布式存储\nIPFS / Pinata"),
    ]
    for x1, y1, x2, y2, txt in boxes:
        d.rounded_rectangle((x1, y1, x2, y2), radius=24, outline="#1f4e79", width=5, fill="#eef6ff")
        lines = txt.split("\n")
        draw_centered(d, (x1, y1 + 18, x2, y1 + 60), lines[0], font)
        draw_centered(d, (x1, y1 + 72, x2, y2 - 12), lines[1], small, fill="#374151")
    for start, end, label in [
        ((360, 185), (500, 185), "HTTP/API"),
        ((780, 185), (920, 185), "ORM"),
        ((640, 250), (640, 420), "ethers.js"),
        ((780, 485), (920, 485), "Metadata URI"),
        ((640, 420), (640, 250), "交易回执"),
    ]:
        d.line((start, end), fill="#0f766e", width=5)
        d.polygon([(end[0], end[1]), (end[0] - 18, end[1] - 10), (end[0] - 18, end[1] + 10)], fill="#0f766e")
        lx = (start[0] + end[0]) / 2
        ly = (start[1] + end[1]) / 2 - 34
        d.text((lx - 70, ly), label, font=small, fill="#0f766e")
    d.text((520, 760), "链上确权与交易，链下查询与展示", font=font, fill="#111827")
    img.save(arch)

    img = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(img)
    steps = [
        "钱包签名登录",
        "上传媒体文件",
        "IPFS生成元数据",
        "合约铸造/上架",
        "市场展示交易",
        "订单同步管理",
    ]
    x = 80
    y = 360
    for i, step in enumerate(steps):
        box = (x + i * 245, y, x + i * 245 + 190, y + 110)
        d.rounded_rectangle(box, radius=18, outline="#7c3aed", width=4, fill="#f5f3ff")
        draw_centered(d, box, step, small, fill="#111827")
        if i < len(steps) - 1:
            sx = box[2]
            ex = x + (i + 1) * 245
            d.line((sx, y + 55, ex, y + 55), fill="#7c3aed", width=5)
            d.polygon([(ex, y + 55), (ex - 18, y + 45), (ex - 18, y + 65)], fill="#7c3aed")
    d.text((430, 170), "NFT数字藏品交易平台核心业务流程", font=font, fill="#111827")
    d.text((280, 640), "主流程强调先完成链上可信操作，再同步链下业务数据。", font=small, fill="#374151")
    img.save(flow)
    return arch, flow


def add_caption(doc: Document, text: str, keep_with_next: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.keep_with_next = keep_with_next
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    r = p.add_run(text)
    set_run_font(r, size=10.5)


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "0")
    tbl_w.set(qn("w:type"), "auto")

    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for name, val, size in [
        ("top", "single", "12"),
        ("left", "nil", "0"),
        ("bottom", "single", "12"),
        ("right", "nil", "0"),
        ("insideH", "single", "4"),
        ("insideV", "nil", "0"),
    ]:
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl_pr.append(borders)


def set_cell_width(cell, cm: float) -> None:
    width = int(cm * 567)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def format_cell(cell, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m in ["top", "left", "bottom", "right"]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), "100")
        node.set(qn("w:type"), "dxa")
    for p in cell.paragraphs:
        p.alignment = align
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(20)
        for run in p.runs:
            set_run_font(run, size=10.5, bold=bold)


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    add_caption(doc, caption, keep_with_next=True)
    table_path = make_table_image(caption, headers, rows, widths)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.keep_together = True
    p.add_run().add_picture(str(table_path), width=Inches(5.9))
    add_paragraph(doc, "", first_line=False, space_after=4)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for ch in text:
        trial = current + ch
        if draw.textlength(trial, font=font) <= max_width or not current:
            current = trial
        else:
            lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def make_table_image(caption: str, headers: list[str], rows: list[list[str]], widths: list[float]) -> Path:
    safe = re.sub(r"[^0-9A-Za-z一-龥-]+", "_", caption).strip("_")
    path = FIG_DIR / f"{safe}.png"
    total_w = 1400
    margin_x = 40
    usable_w = total_w - margin_x * 2
    width_sum = sum(widths)
    col_px = [int(usable_w * w / width_sum) for w in widths]
    col_px[-1] = usable_w - sum(col_px[:-1])
    font = get_font(30)
    header_font = get_font(31)
    pad_x = 18
    pad_y = 16

    tmp = Image.new("RGB", (total_w, 400), "white")
    dtmp = ImageDraw.Draw(tmp)

    all_rows = [headers] + rows
    row_heights: list[int] = []
    wrapped: list[list[list[str]]] = []
    for ridx, row in enumerate(all_rows):
        row_lines: list[list[str]] = []
        max_lines = 1
        f = header_font if ridx == 0 else font
        for cidx, cell in enumerate(row):
            lines = wrap_text(dtmp, cell, f, col_px[cidx] - pad_x * 2)
            row_lines.append(lines)
            max_lines = max(max_lines, len(lines))
        wrapped.append(row_lines)
        row_heights.append(max(72, max_lines * 38 + pad_y * 2))

    total_h = sum(row_heights) + 60
    img = Image.new("RGB", (total_w, total_h), "white")
    draw = ImageDraw.Draw(img)
    left = margin_x
    top = 28
    right = total_w - margin_x
    y = top
    draw.line((left, y, right, y), fill="black", width=4)
    for ridx, row_lines in enumerate(wrapped):
        height = row_heights[ridx]
        x = left
        for cidx, lines in enumerate(row_lines):
            cell_w = col_px[cidx]
            f = header_font if ridx == 0 else font
            line_h = 38
            block_h = len(lines) * line_h
            text_y = y + (height - block_h) / 2 - 2
            if cidx == len(row_lines) - 1 and ridx != 0:
                text_x = x + pad_x
                for line in lines:
                    draw.text((text_x, text_y), line, font=f, fill="#111111")
                    text_y += line_h
            else:
                for line in lines:
                    line_w = draw.textlength(line, font=f)
                    draw.text((x + (cell_w - line_w) / 2, text_y), line, font=f, fill="#111111")
                    text_y += line_h
            x += cell_w
        y += height
        if ridx == 0:
            draw.line((left, y, right, y), fill="black", width=3)
    draw.line((left, y, right, y), fill="black", width=4)
    img.save(path)
    return path


CHAPTERS = [
    {
        "title": "第1章 绪论",
        "sections": [
            ("1.1 选题背景", [
                "在数字经济持续发展的背景下，数字内容已经成为互联网平台的重要组成部分。图片、音乐、短视频、数字插画和虚拟纪念品等内容可以被快速创建和传播，但传统文件本身很难证明唯一性，也难以完整记录作品从发布到交易的流转过程。对于创作者而言，作品被复制、二次传播甚至被他人包装销售的情况并不少见；对于购买者而言，仅凭平台数据库记录也难以确认某个数字藏品是否具有可验证的归属凭证。",
                "区块链技术为数字资产确权提供了新的思路。通过密码学签名、链上账本和智能合约，数字作品可以被映射为具有唯一编号的 NFT，并在链上记录所有权变化。ERC-721 标准规定了非同质化通证的基本接口，使不同钱包、应用和区块浏览器能够识别 NFT 的归属关系[1]。在这一基础上，数字藏品交易平台可以把作品展示、市场筛选、订单管理和用户资料等功能放在链下系统中，把确权和交易等关键环节放在链上执行。",
            ]),
            ("1.2 研究目的与意义", [
                "本文的研究目的，是围绕 NFT 数字藏品交易场景，设计并实现一个具备完整业务闭环的平台系统。系统不仅要实现 NFT 铸造，还要支持钱包签名登录、媒体文件上传、IPFS 元数据存储、市场浏览、链上购买、订单同步、个人中心和通知提醒等功能。换句话说，本课题关注的不是单一合约演示，而是一个用户真正能操作的数字藏品市场原型。",
                "从理论意义看，本课题能够说明链上链下协同架构在 NFT 应用中的边界划分。链上适合承担资产确权、购买付款和版税结算，链下适合承担高频查询、用户资料和页面展示。二者协同，可以在可信性和可用性之间取得平衡。从实践意义看，本系统综合应用前端开发、后端服务、数据库建模、智能合约、IPFS 存储和钱包交互，能够体现软件设计类毕业论文对工程完整性的要求。",
            ]),
            ("1.3 主要研究内容", [
                "本文围绕项目源码和功能实现展开研究。首先，对 NFT 数字藏品平台的业务需求进行分析，明确用户、藏品、订单、媒体存储和交易同步等核心对象。其次，设计前端、后端、数据库、智能合约和 IPFS 的总体架构。再次，完成钱包认证、NFT 创建、市场展示、交易购买、个人中心和通知提醒等模块实现。最后，通过功能测试和源码验证分析系统实现效果，并总结系统特点、存在问题和后续改进方向。",
                "论文正文按照软件设计类毕业论文常见结构组织，包括绪论、相关开发技术、系统分析、系统设计、系统实现、系统测试和结论七个部分。其中结论独立成章，不再使用章节编号，重点归纳系统成果、特色和不足。",
            ]),
        ],
    },
    {
        "title": "第2章 相关开发技术",
        "sections": [
            ("2.1 区块链、NFT与智能合约", [
                "区块链是一种由多节点共同维护的分布式账本，具有数据难篡改、交易可追溯和状态公开验证等特点。NFT 是运行在区块链上的非同质化数字资产，每个 Token 都具有独立编号和元数据，不同 Token 之间不能简单等价互换。ERC-721 是以太坊生态中常见的 NFT 标准，定义了所有权查询、转移和元数据访问等接口[1]。本项目中的 NFTCollection 合约正是围绕 ERC-721 扩展实现。",
                "智能合约是部署在区块链上的程序逻辑。与普通后端程序相比，智能合约一旦部署后具有更强的公开性和可验证性，但也受到执行成本和修改困难等限制。因此，本系统把所有权转移、上架状态、购买付款和版税分配放在合约中执行，而把市场搜索、列表展示、个人资料和订单聚合放在后端数据库中处理。",
            ]),
            ("2.2 EIP-2981版税机制", [
                "数字藏品交易中，创作者往往希望作品在二级交易时仍能获得一定收益。EIP-2981 提供了 NFT 版税信息标准，允许合约根据销售价格返回版税接收地址和应支付金额[2]。本系统在合约中继承 ERC721Royalty，并设置最大版税比例为 2500 bps，即 25%。用户创建 NFT 时可以设置版税接收地址和比例，购买时由合约完成版税与卖家实收金额的拆分。",
                "OpenZeppelin Contracts 为 ERC-721、ERC-2981、Ownable 和 ReentrancyGuard 等常用合约能力提供了成熟实现[3]。项目复用这些基础组件，减少了从零实现标准接口和安全保护的复杂度，也让合约结构更容易解释和维护。",
            ]),
            ("2.3 IPFS与Pinata存储", [
                "NFT 的媒体文件通常不直接写入链上，因为链上存储成本高且不适合大体积文件。IPFS 使用内容寻址方式，通过 CID 标识文件内容，使资源定位不再依赖单一服务器路径[5]。本项目通过 Pinata 提供的文件上传能力接入 IPFS，后端将主媒体文件、封面文件和元数据 JSON 上传后返回 CID、ipfs URI 和网关访问地址[6]。",
                "在本系统中，图片类 NFT 使用元数据中的 image 字段指向资源；音频和视频类 NFT 使用 animation_url 字段，并可额外设置封面图。这样的元数据组织方式能够兼顾钱包、市场页面和浏览器访问需要。",
            ]),
            ("2.4 前端与后端技术", [
                "前端采用 Next.js 14 与 React 18。Next.js 的 App Router 适合组织首页、市场页、详情页、创建页和个人中心等多页面应用[7]，React 组件化机制则便于拆分卡片、分类标签、通知、钱包按钮和交易进度组件。ethers.js v6 用于连接浏览器钱包、构造合约对象、发起交易和解析交易回执[8]。",
                "后端采用 Go 1.22、Gorilla Mux、GORM 和 MySQL。Go 适合构建轻量、可部署的后端服务[9]；Gorilla Mux 用于 REST 路由组织[10]；GORM 用于对象关系映射和数据库迁移[11]；MySQL 负责用户、NFT 和订单数据的持久化存储[12]。JWT 用于后端会话鉴权，其标准形式由 RFC 7519 定义[13]。",
            ]),
        ],
    },
    {
        "title": "第3章 系统分析",
        "sections": [
            ("3.1 用户角色分析", [
                "本系统没有把用户严格划分为互斥角色，而是采用统一用户模型。一个完成钱包登录的用户，可以作为创作者发布作品，也可以作为卖家上架自己持有的 NFT，还可以作为买家购买其他用户发布的 NFT。这种设计更符合数字藏品市场的实际使用方式，因为钱包地址本身既可以持有资产，也可以发起交易。",
                "用户的核心行为包括：连接钱包并签名登录，上传作品并创建 NFT，浏览和筛选市场藏品，在详情页发起购买，管理个人资料，查看当前持有和历史交易。平台的业务目标就是让这些行为形成连续流程，而不是彼此割裂。",
            ]),
            ("3.2 功能需求分析", [
                "根据项目实现和业务目标，系统功能需求可划分为身份认证、NFT 创建、市场浏览、详情交易、订单管理、个人中心和通知提醒等模块。各模块既有独立职责，也存在数据依赖。例如 NFT 创建依赖 IPFS 上传和合约铸造，购买流程依赖钱包交易和后端订单同步，个人中心依赖 NFT 与订单数据聚合。",
            ], "functional_table"),
            ("3.3 非功能需求分析", [
                "除基本功能外，系统还需要满足安全性、一致性、可用性和可扩展性要求。安全性要求包括签名登录、JWT 鉴权、文件类型校验、上传大小限制和合约重入保护；一致性要求主要体现在链上交易完成后，链下数据库能够正确更新 NFT 拥有者和订单记录；可用性要求体现在交易进度、错误提示、媒体预览和响应式布局；可扩展性要求体现在模块边界清晰，后续可以增加事件监听、后台管理和多链支持。",
            ], "nonfunctional_table"),
            ("3.4 可行性分析", [
                "技术可行性方面，项目采用的 Next.js、Go、MySQL、Solidity、Hardhat、IPFS 等技术均有成熟生态和文档支持。经济可行性方面，系统主要依赖开源框架和测试网络环境，开发成本较低。操作可行性方面，用户只需安装浏览器钱包，即可完成登录、创建和交易流程。虽然链上交易需要等待确认，但前端进度组件能够降低用户理解成本。",
            ]),
        ],
    },
    {
        "title": "第4章 系统设计",
        "sections": [
            ("4.1 总体架构设计", [
                "系统采用前后端分离与链上链下协同架构。前端负责用户界面、钱包交互和交易反馈；后端负责 REST API、认证、文件校验、IPFS 上传和数据库管理；智能合约负责 NFT 铸造、上架、购买、下架和版税结算；MySQL 保存用户、NFT 和订单等链下业务数据；IPFS 保存媒体文件和元数据。总体架构如图 4-1 所示。",
            ], "architecture_figure"),
            ("4.2 数据库设计", [
                "数据库设计围绕 users、nfts 和 orders 三张核心表展开。users 表保存平台用户和钱包登录所需 nonce，nfts 表保存链上 NFT 在平台中的展示和交易状态，orders 表保存历史成交记录。三者之间的关系为：一个用户可以拥有多个 NFT，一个用户可以作为买家或卖家参与多笔订单，一个 NFT 可以产生多笔历史交易记录。",
            ], "database_tables"),
            ("4.3 智能合约设计", [
                "NFTCollection 合约基于 ERC721URIStorage、ERC721Royalty、Ownable 和 ReentrancyGuard 组合实现。合约内部通过自增计数器生成 Token ID，通过 Listing 结构保存上架卖家、价格和状态，通过版税映射保存每个 Token 的版税基点。合约暴露 safeMint、mintAndList、listToken、cancelListing、buy、getListing、getRoyaltyInfo 和 totalMinted 等接口。",
                "购买函数是合约的关键。它先检查 NFT 是否处于有效上架状态，再检查卖家是否仍为链上持有人、买家是否不是卖家、付款金额是否等于价格。随后合约删除上架信息、转移 NFT 所有权，并根据版税信息拆分支付金额。该流程保证了链上交易的可信执行。",
            ]),
            ("4.4 核心业务流程设计", [
                "系统核心业务流程包括钱包登录、媒体上传、IPFS 元数据生成、合约铸造或上架、市场展示交易和订单同步管理。为了保持链上链下状态一致，系统采用“链上交易确认后，再同步后端业务数据”的设计。核心流程如图 4-2 所示。",
            ], "flow_figure"),
        ],
    },
    {
        "title": "第5章 系统实现",
        "sections": [
            ("5.1 项目结构实现", [
                "项目根目录主要包含 frontend、backend、contracts 和 hardhat 四个模块。frontend 使用 Next.js App Router 组织页面，backend 使用 Go 模块组织服务端代码，contracts 保存合约源码，hardhat 用于合约编译和部署。当前一方源码统计约 7823 行，其中前端约 5498 行，后端约 1943 行，合约约 175 行，Hardhat 配置与部署相关代码约 207 行。",
            ]),
            ("5.2 身份认证模块实现", [
                "身份认证由后端 AuthService 和前端 WalletConnectButton 协同完成。后端首先根据钱包地址生成随机 nonce，并在用户不存在时自动创建用户；前端调用钱包对固定格式消息签名；后端恢复签名地址并与请求地址比对，成功后签发 JWT。登录成功后，前端把 token 和用户信息保存到 localStorage，并通过 Axios 拦截器在后续 API 请求中自动携带 Authorization 头。",
                "该实现避免了账号密码体系与链上钱包身份脱节的问题，也降低了密码泄露风险。用户第一次连接钱包时，系统就能自动完成建号和登录。",
            ]),
            ("5.3 NFT创建与IPFS模块实现", [
                "创建 NFT 时，前端页面负责收集名称、描述、分类、价格、版税和文件信息，并根据分类限制上传类型。后端 IPFS 接口负责文件大小和 MIME 类型校验，主文件最大 100MB，封面最大 20MB。文件通过 Pinata 上传到 IPFS 后，后端自动生成元数据 JSON，再将元数据上传到 IPFS。前端拿到 metadataUri 后调用合约完成铸造，并将 tokenId、txHash、媒体地址和价格信息同步到数据库。",
            ]),
            ("5.4 市场交易模块实现", [
                "市场页通过 NFT 列表接口获取已上架藏品，并支持分类筛选、关键词搜索、价格区间和排序。详情页根据 NFT 所有者判断展示购买按钮或上架管理按钮。普通买家购买时，前端先读取链上 listing，再调用 buy 函数完成付款和所有权转移。交易确认后，前端提交 txHash 到后端，后端在事务中创建订单并更新 NFT owner。",
            ], "api_table"),
            ("5.5 个人中心与通知模块实现", [
                "个人中心展示用户头像、用户名、钱包地址、当前持有 NFT、历史买入和历史卖出记录。头像上传使用本地 uploads 目录，NFT 媒体则使用 IPFS，二者根据资源性质采用不同存储方式。通知模块通过定时轮询买入和卖出订单接口，检测新增订单并生成本地未读提醒。虽然该方式不是严格实时推送，但能满足毕业设计阶段对成交提醒的展示需求。",
            ]),
        ],
    },
    {
        "title": "第6章 系统测试",
        "sections": [
            ("6.1 测试环境", [
                "系统测试以源码阅读、功能流程验证和后端编译检查为主。后端在 backend 目录执行 go test ./...，结果显示各包均可通过检查，但当前暂无独立测试文件。前端和合约流程主要通过页面操作、钱包交互和接口调用进行联调验证。",
            ], "env_table"),
            ("6.2 功能测试", [
                "功能测试围绕主业务闭环设计，包括钱包登录、NFT 创建、IPFS 上传、市场浏览、NFT 购买、订单同步、个人中心和通知提醒等场景。测试重点不是单个按钮是否能点击，而是链路是否能从用户操作顺利到达数据库和链上状态变化。",
            ], "case_table"),
            ("6.3 测试结果分析", [
                "从测试结果看，系统已经完成主要业务闭环。用户能够通过钱包签名登录，能够上传媒体文件并创建 NFT，能够在市场页浏览已上架藏品，能够在详情页发起真实链上购买，购买完成后订单和 NFT 当前拥有者能够同步更新。系统仍存在自动化测试不足的问题，后续可补充 Go 单元测试、接口测试、Hardhat 合约测试和 Playwright 前端端到端测试。",
            ]),
        ],
    },
]


def add_special(doc: Document, marker: str, images: tuple[Path, Path]) -> None:
    if marker == "functional_table":
        add_table(
            doc,
            "表 3-1 功能需求表",
            ["模块", "主要功能", "说明"],
            [
                ["身份认证", "钱包连接、nonce 签名、JWT", "证明用户控制对应钱包地址"],
                ["NFT创建", "上传媒体、设置价格、设置版税", "完成作品从文件到链上资产的转换"],
                ["市场浏览", "分类、搜索、排序、收藏", "提高用户发现目标藏品的效率"],
                ["详情交易", "购买、上架、下架、交易记录", "承担主要交易入口"],
                ["个人中心", "资料维护、持有资产、历史订单", "展示用户长期资产状态"],
            ],
            [2.8, 4.5, 7.0],
        )
    elif marker == "nonfunctional_table":
        add_table(
            doc,
            "表 3-2 非功能需求表",
            ["类别", "需求内容", "实现体现"],
            [
                ["安全性", "身份、防刷、文件和交易安全", "签名登录、JWT、限流、MIME 校验、合约重入保护"],
                ["一致性", "链上交易与链下数据同步", "交易确认后创建订单并更新 owner"],
                ["可用性", "降低用户操作理解成本", "交易进度、媒体预览和错误提示"],
                ["可扩展性", "后续支持更多能力", "模块化目录、标准合约接口和事件设计"],
            ],
            [2.8, 4.8, 6.7],
        )
    elif marker == "architecture_figure":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.add_run().add_picture(str(images[0]), width=Inches(5.9))
        add_caption(doc, "图 4-1 系统总体架构图")
    elif marker == "flow_figure":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.add_run().add_picture(str(images[1]), width=Inches(5.9))
        add_caption(doc, "图 4-2 核心业务流程图")
        add_paragraph(
            doc,
            "从图 4-2 可以看出，平台流程并不是简单的页面跳转，而是由链下业务校验、IPFS 元数据保存、链上合约交易和后端订单同步共同构成。前端负责把用户操作组织成连续步骤，后端负责保存可查询的数据状态，智能合约负责完成不可篡改的资产流转。这样的流程安排能够降低单点数据失真的风险，也方便后续通过事件监听进一步提高链上链下同步的自动化程度。",
        )
    elif marker == "database_tables":
        doc.add_page_break()
        add_table(
            doc,
            "表 4-1 用户表 users 主要字段",
            ["字段", "类型", "说明"],
            [
                ["id", "uint", "用户主键"],
                ["wallet", "string", "钱包地址，唯一索引"],
                ["username", "string", "平台展示用户名"],
                ["avatar", "string", "用户头像地址"],
                ["nonce", "string", "钱包签名登录的一次性随机数"],
            ],
            [3.0, 3.0, 7.9],
        )
        add_table(
            doc,
            "表 4-2 NFT 表 nfts 主要字段",
            ["字段", "类型", "说明"],
            [
                ["contract", "string", "链上合约地址"],
                ["token_id", "string", "链上 Token ID"],
                ["media_url", "string", "主媒体文件访问地址"],
                ["token_uri", "string", "链上元数据 URI"],
                ["price_wei", "string", "链上精确价格，单位 Wei"],
                ["royalty_fee_bps", "uint16", "版税比例，单位 bps"],
            ],
            [3.2, 3.0, 7.7],
        )
        doc.add_page_break()
        add_table(
            doc,
            "表 4-3 订单表 orders 主要字段",
            ["字段", "类型", "说明"],
            [
                ["nft_id", "uint", "被交易 NFT 编号"],
                ["buyer_id", "uint", "买家用户编号"],
                ["seller_id", "uint", "卖家用户编号"],
                ["price_wei", "string", "成交价格，单位 Wei"],
                ["tx_hash", "string", "链上交易哈希"],
                ["status", "string", "订单状态"],
            ],
            [3.2, 3.0, 7.7],
        )
    elif marker == "api_table":
        add_table(
            doc,
            "表 5-1 后端主要接口表",
            ["接口", "方法", "功能说明"],
            [
                ["/auth/wallet/nonce", "GET", "生成钱包登录 nonce"],
                ["/auth/wallet/login", "POST", "校验签名并签发 JWT"],
                ["/ipfs/nft", "POST", "上传媒体并生成 IPFS 元数据"],
                ["/nfts", "POST/GET", "创建或查询 NFT"],
                ["/orders", "POST", "创建成交订单并更新归属"],
                ["/orders/bought", "GET", "查询当前用户买入记录"],
            ],
            [4.0, 2.5, 7.4],
        )
    elif marker == "env_table":
        add_table(
            doc,
            "表 6-1 测试环境表",
            ["项目", "环境或版本", "说明"],
            [
                ["前端", "Next.js 14 / React 18", "页面与组件测试环境"],
                ["后端", "Go 1.22", "REST API 与业务逻辑"],
                ["数据库", "MySQL / GORM", "用户、NFT、订单数据"],
                ["合约", "Solidity 0.8.24 / Hardhat", "NFTCollection 编译与部署"],
                ["钱包", "MetaMask / OKX / Bitget", "浏览器注入钱包交互"],
            ],
            [3.0, 4.5, 6.4],
        )
    elif marker == "case_table":
        doc.add_page_break()
        add_table(
            doc,
            "表 6-2 功能测试用例表",
            ["编号", "测试项", "预期结果", "结论"],
            [
                ["1", "钱包签名登录", "返回 JWT 并保存用户信息", "通过"],
                ["2", "NFT 文件上传", "返回 assetCid 和 metadataUri", "通过"],
                ["3", "NFT 铸造上架", "链上返回 tokenId 和 txHash", "通过"],
                ["4", "市场筛选搜索", "列表按条件刷新", "通过"],
                ["5", "NFT 购买", "链上转移所有权并生成订单", "通过"],
                ["6", "个人中心查询", "展示持有和历史订单", "通过"],
            ],
            [1.8, 3.5, 5.4, 2.2],
        )


def add_body(doc: Document, images: tuple[Path, Path]) -> None:
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_page_layout(sec)
    restart_page_numbering(sec, 1)
    setup_body_header_footer(sec)

    first = True
    for chapter in CHAPTERS:
        if not first:
            doc.add_page_break()
        first = False
        add_heading(doc, chapter["title"], 1)
        for item in chapter["sections"]:
            title, paragraphs = item[0], item[1]
            marker = item[2] if len(item) > 2 else None
            add_heading(doc, title, 2)
            for para in paragraphs:
                add_paragraph(doc, para)
            if marker:
                add_special(doc, marker, images)

    doc.add_page_break()
    add_heading(doc, "结论", 1)
    conclusion = [
        "本文围绕 NFT 数字藏品交易场景，完成了一套基于区块链的交易平台设计与实现。系统以 NFTCollection 智能合约作为链上核心，以 Go 后端作为链下业务服务，以 Next.js 前端作为用户交互入口，并通过 Pinata 接入 IPFS 完成媒体文件和元数据存储。系统实现了钱包签名登录、NFT 多媒体创建、链上铸造、上架、购买、下架、版税结算、订单同步、市场浏览、个人中心和通知提醒等功能，能够支撑数字藏品从发布到交易再到资产管理的完整流程。",
        "从设计结果看，系统的主要特点体现在三个方面。第一，系统采用链上链下协同架构，将确权、支付和版税等关键逻辑放在智能合约中执行，将高频查询、页面展示和用户资料管理放在后端数据库中处理。第二，系统支持图片、音频和视频等多媒体 NFT，并通过 IPFS 元数据组织提升资源可验证性。第三，系统在前端加入交易进度、媒体预览和个人资产视图，使用户能够更清楚地理解 Web3 操作过程。",
        "当前系统仍存在一些不足，例如链上链下同步主要依赖前端回传，后端尚未实现合约事件监听；自动化测试覆盖不足，主要依赖功能联调和编译检查；部分配置仍偏开发环境。后续可以围绕事件监听、测试体系、生产配置和多链适配继续改进。总体而言，本课题较好地完成了预期目标，体现了区块链、NFT、智能合约、IPFS、前后端工程和数据库设计的综合应用能力。",
    ]
    for para in conclusion:
        add_paragraph(doc, para)

    add_references(doc)
    add_appendix(doc)
    add_acknowledgement(doc)


def add_references(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "参考文献", 1)
    refs = [
        "Ethereum Improvement Proposals. ERC-721: Non-Fungible Token Standard[EB/OL]. https://eips.ethereum.org/EIPS/eip-721, 2018-01-24/2026-04-25.",
        "Ethereum Improvement Proposals. ERC-2981: NFT Royalty Standard[EB/OL]. https://eips.ethereum.org/EIPS/eip-2981, 2020-09-15/2026-04-25.",
        "OpenZeppelin. OpenZeppelin Contracts 5.x Documentation[EB/OL]. https://docs.openzeppelin.com/contracts/5.x, 2026-04-25.",
        "Solidity Team. Solidity 0.8.24 Release Announcement[EB/OL]. https://soliditylang.org/blog/2024/01/26/solidity-0.8.24-release-announcement/, 2024-01-26/2026-04-25.",
        "IPFS Docs. Content Identifiers (CIDs)[EB/OL]. https://docs.ipfs.tech/concepts/content-addressing/, 2026-04-25.",
        "Pinata. Uploading Files - Pinata Docs[EB/OL]. https://docs.pinata.cloud/files/uploading-files, 2026-04-25.",
        "Vercel. Next.js 14 App Router Documentation[EB/OL]. https://nextjs.org/docs/14/app, 2026-04-25.",
        "ethers.org. ethers.js v6 Documentation[EB/OL]. https://docs.ethers.org/v6/, 2026-04-25.",
        "The Go Authors. Go 1.22 Release Notes[EB/OL]. https://go.dev/doc/go1.22, 2024-02-06/2026-04-25.",
        "Gorilla. gorilla/mux: HTTP Router and URL Matcher[EB/OL]. https://github.com/gorilla/mux, 2026-04-25.",
        "GORM. GORM Guides[EB/OL]. https://gorm.io/docs/, 2026-04-25.",
        "Oracle. MySQL 8.4 Reference Manual[EB/OL]. https://dev.mysql.com/doc/refman/en/, 2026-04-25.",
        "Jones M, Bradley J, Sakimura N. RFC 7519: JSON Web Token (JWT)[S/OL]. https://www.rfc-editor.org/rfc/rfc7519, 2015-05/2026-04-25.",
    ]
    for idx, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(21)
        p.paragraph_format.first_line_indent = Pt(-21)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(20)
        r = p.add_run(f"[{idx}] {ref}")
        set_run_font(r, size=10.5)


def add_appendix(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "附录", 1)
    add_heading(doc, "附录A 系统主要路由", 2)
    routes = [
        "GET /api/v1/auth/wallet/nonce：根据钱包地址生成登录 nonce。",
        "POST /api/v1/auth/wallet/login：校验签名并返回 JWT。",
        "POST /api/v1/ipfs/nft：上传 NFT 媒体文件并生成 IPFS 元数据。",
        "GET /api/v1/nfts：获取 NFT 列表，支持分类和上架筛选。",
        "POST /api/v1/orders：根据链上交易哈希创建订单并同步资产归属。",
    ]
    for item in routes:
        add_paragraph(doc, item)
    add_heading(doc, "附录B 智能合约主要事件", 2)
    events = [
        "Minted：记录 NFT 铸造接收者、Token ID 和 tokenURI。",
        "Listed：记录 NFT 上架的卖家地址和价格。",
        "Delisted：记录 NFT 下架行为。",
        "RoyaltySet：记录指定 Token 的版税接收地址和版税比例。",
        "Purchased：记录买家、卖家、成交价格、版税金额和卖家实收金额。",
    ]
    for item in events:
        add_paragraph(doc, item)


def add_acknowledgement(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "致谢", 1)
    paras = [
        "在本课题完成过程中，指导教师在选题方向、系统设计和论文结构方面给予了重要帮助，使我能够围绕 NFT 数字藏品交易平台这一主题逐步完成需求分析、系统设计、功能实现和论文整理工作。",
        "同时，感谢同学和朋友在系统联调、页面体验和问题排查中提供的建议。通过本次毕业设计，我对区块链应用开发、前后端协同、数据库建模、智能合约和 IPFS 存储有了更加系统的理解。谨向所有给予帮助和支持的老师、同学和家人表示诚挚感谢。",
    ]
    for para in paras:
        add_paragraph(doc, para)


def collect_plain_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


def build() -> None:
    ensure_dirs()
    images = make_diagrams()
    doc = Document()
    set_doc_defaults(doc)
    add_cover(doc)
    add_abstract(doc)
    add_toc(doc)
    add_body(doc, images)
    doc.save(DOCX_PATH)

    text = collect_plain_text(doc)
    summary = {
        "title": TITLE,
        "title_length_without_ascii": len(re.sub(r"[A-Za-z0-9]", "", TITLE)) + len(re.findall(r"[A-Za-z0-9]+", TITLE)),
        "abstract_chars": sum(len(p) for p in abstract_text()),
        "body_total_chars": len(text),
        "output": str(DOCX_PATH),
    }
    (RAW_DIR / "thesis_content.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    build()
